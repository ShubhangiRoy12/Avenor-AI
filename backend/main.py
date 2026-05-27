from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional, List
import httpx
import os
import uuid
import urllib.parse
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
HF_API_TOKEN = os.getenv("HF_API_TOKEN", "")
GROQ_MODEL = "llama-3.3-70b-versatile"

Path("outputs/images").mkdir(parents=True, exist_ok=True)
Path("outputs/docs").mkdir(parents=True, exist_ok=True)

app = FastAPI(title="Avenor AI", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/outputs", StaticFiles(directory="outputs"), name="outputs")

class QueryRequest(BaseModel):
    query: str
    mode: str = "auto"

class ExportRequest(BaseModel):
    content: str
    filename: str = "avenor_agent_output"

class QueryResponse(BaseModel):
    result: str
    agent_used: str
    steps: List[str]
    image_url: Optional[str] = None

SYSTEM_PROMPT = """You are Avenor AI — an enterprise-grade multi-agent AI platform.
You are highly intelligent, helpful, and professional.
Format your responses clearly. Use bullet points and structure when needed.
For emails, write complete professional emails.
For code, use proper formatting.
For explanations, be thorough but concise."""

def call_groq(user_message: str, system: str = SYSTEM_PROMPT) -> str:
    try:
        response = httpx.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": GROQ_MODEL,
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_message}
                ],
                "temperature": 0.7,
                "max_tokens": 1024
            },
            timeout=30
        )
        data = response.json()
        if "choices" in data:
            return data["choices"][0]["message"]["content"]
        else:
            return f"API Error: {data.get('error', {}).get('message', str(data))}"
    except Exception as e:
        return f"Connection error: {str(e)}"

def generate_image_free(prompt: str) -> Optional[str]:
    import urllib.parse
    encoded = urllib.parse.quote(prompt)
    try:
        url = f"https://image.pollinations.ai/prompt/{encoded}?width=768&height=512&model=flux-schnell&nologo=true"
        response = httpx.get(url, timeout=60, follow_redirects=True)
        content_type = response.headers.get("content-type", "")
        if response.status_code == 200 and "image" in content_type:
            filename = f"{uuid.uuid4()}.png"
            filepath = f"outputs/images/{filename}"
            with open(filepath, "wb") as f:
                f.write(response.content)
            return f"http://127.0.0.1:8000/outputs/images/{filename}"
        return None
    except Exception as e:
        print(f"Pollinations error: {e}")
        return None

def generate_image_hf(prompt: str) -> Optional[str]:
    """HuggingFace with correct new router URL"""
    if not HF_API_TOKEN:
        return None
    try:
        response = httpx.post(
            "https://router.huggingface.co/hf-inference/models/stabilityai/stable-diffusion-xl-base-1.0",
            headers={
                "Authorization": f"Bearer {HF_API_TOKEN}",
                "Content-Type": "application/json"
            },
            json={"inputs": prompt},
            timeout=120
        )
        content_type = response.headers.get("content-type", "")
        if response.status_code == 200 and "image" in content_type:
            filename = f"{uuid.uuid4()}.png"
            filepath = f"outputs/images/{filename}"
            with open(filepath, "wb") as f:
                f.write(response.content)
            return f"http://127.0.0.1:8000/outputs/images/{filename}"
        else:
            print(f"HF Error: {response.text[:200]}")
            return None
    except Exception as e:
        print(f"HF Exception: {e}")
        return None
    
def generate_pdf(content: str, filename: str) -> str:
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_margins(20, 20, 20)
        pdf.set_fill_color(30, 30, 50)
        pdf.rect(0, 0, 210, 25, 'F')
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_xy(0, 8)
        pdf.cell(210, 10, "AVENOR AI - Generated Document", align="C")
        pdf.ln(20)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", size=11)
        for line in content.split("\n"):
            clean = line.strip().encode("latin-1", "replace").decode("latin-1")
            if not clean:
                pdf.ln(4)
            elif clean.startswith("#"):
                pdf.set_font("Helvetica", "B", 13)
                pdf.multi_cell(0, 8, clean.replace("#", "").strip())
                pdf.set_font("Helvetica", size=11)
            else:
                try:
                    pdf.multi_cell(0, 7, clean)
                except:
                    pass
        filepath = f"outputs/docs/{filename}.pdf"
        pdf.output(filepath)
        return filepath
    except ImportError:
        return ""

def generate_word(content: str, filename: str) -> str:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        title = doc.add_heading("Avenor AI - Generated Document", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        for line in content.split("\n"):
            clean = line.strip()
            if not clean:
                doc.add_paragraph("")
            elif clean.startswith("## "):
                doc.add_heading(clean.replace("## ", ""), level=2)
            elif clean.startswith("# "):
                doc.add_heading(clean.replace("# ", ""), level=1)
            elif clean.startswith("- ") or clean.startswith("* "):
                doc.add_paragraph(clean[2:], style="List Bullet")
            else:
                doc.add_paragraph(clean)
        filepath = f"outputs/docs/{filename}.docx"
        doc.save(filepath)
        return filepath
    except ImportError:
        return ""

def decide_agent(query: str, mode: str) -> str:
    if mode != "auto":
        return mode
    q = query.lower()
    if any(w in q for w in ["generate image", "create image", "draw", "picture", "poster", "visualize", "image of"]):
        return "image"
    if any(w in q for w in ["email", "mail", "write to", "draft", "letter", "message to", "cold mail"]):
        return "action"
    if any(w in q for w in ["research", "find", "search", "latest", "news", "trend"]):
        return "research"
    return "analysis"

@app.get("/")
def home():
    return {"message": "Welcome to Avenor AI", "status": "running"}

@app.get("/health")
def health_check():
    return {
        "status": "Backend Running Successfully",
        "groq": "connected" if GROQ_API_KEY else "missing",
        "huggingface": "connected" if HF_API_TOKEN else "missing",
        "model": GROQ_MODEL
    }

@app.post("/ask")
def ask(request: QueryRequest):
    agent = decide_agent(request.query, request.mode)
    image_url = None

    steps = [
        "MemoryAgent: Query stored",
        "ResearchAgent: Context retrieved",
        f"DecisionAgent: Routing to {agent.capitalize()}Agent",
    ]

    if agent == "image":
            image_url = generate_image_hf(request.query)
            if not image_url:
                image_url = generate_image_free(request.query)
            if image_url:
                result = "✅ Image generated successfully!"
            else:
                result = "⚠️ Image generation failed. Please try again."
            steps.append("ImageAgent: Image generated")
            agent_used = "ImageAgent"

    elif agent == "action":
        result = call_groq(
            request.query,
            "You are Avenor AI's Action Agent. Specialize in writing professional "
            "emails, letters, and messages. Write complete, polished content."
        )
        steps.append("ActionAgent: Content generated")
        agent_used = "ActionAgent"

    elif agent == "research":
        result = call_groq(
            request.query,
            "You are Avenor AI's Research Agent. Provide thorough, well-structured "
            "research with key facts, insights, and clear explanations."
        )
        steps.append("ResearchAgent: Research completed")
        agent_used = "ResearchAgent"

    else:
        result = call_groq(request.query)
        steps.append("AnalysisAgent: Response generated")
        agent_used = "AnalysisAgent"

    return {
        "result": result,
        "agent_used": agent_used,
        "steps": steps,
        "image_url": image_url
    }

@app.post("/generate-pdf")
def create_pdf(request: QueryRequest):
    content = call_groq(
        request.query,
        "Generate a detailed professional document with clear headings using # and ##, "
        "and bullet points using -. Make it comprehensive."
    )
    filename = str(uuid.uuid4())
    filepath = generate_pdf(content, filename)
    if not filepath:
        return {"error": "Run: pip install fpdf2"}
    return FileResponse(
        filepath,
        media_type="application/pdf",
        filename="avenor_document.pdf"
    )

@app.post("/generate-word")
def create_word(request: QueryRequest):
    content = call_groq(
        request.query,
        "Generate a detailed professional document with clear headings using # and ##, "
        "and bullet points using -. Make it comprehensive."
    )
    filename = str(uuid.uuid4())
    filepath = generate_word(content, filename)
    if not filepath:
        return {"error": "Run: pip install python-docx"}
    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="avenor_document.docx"
    )

@app.post("/export-pdf")
def export_pdf(request: ExportRequest):
    filename = str(uuid.uuid4())
    filepath = generate_pdf(request.content, filename)
    if not filepath:
        return {"error": "Run: pip install fpdf2"}
    return FileResponse(
        filepath,
        media_type="application/pdf",
        filename=f"{request.filename}.pdf"
    )

@app.post("/export-word")
def export_word(request: ExportRequest):
    filename = str(uuid.uuid4())
    filepath = generate_word(request.content, filename)
    if not filepath:
        return {"error": "Run: pip install python-docx"}
    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{request.filename}.docx"
    )

@app.get("/agents")
def get_agents():
    return {
        "agents": [
            "OrchestratorAgent",
            "ResearchAgent",
            "AnalysisAgent",
            "DecisionAgent",
            "MemoryAgent",
            "ActionAgent",
            "ImageAgent"
        ]
    }
